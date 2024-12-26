from flask import Flask, request, jsonify, render_template, send_file, after_this_request
import vertexai
from vertexai.generative_models import GenerativeModel
from docx import Document
import os
import re
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

# Initialize Vertex AI
vertexai.init(project="bamboo-basis-434518-f9", location="us-central1")

# Function to sanitize file names
def sanitize_filename(domain):
    return re.sub(r'[^\w\-_\. ]', '_', domain)  # Replace non-alphanumeric characters with underscores

# Serve the HTML file from the 'templates' directory
@app.route('/')
def index():
    return render_template('index.html')

# Endpoint to get internship tasks as Markdown content
@app.route('/generate-internship', methods=['POST'])
def generate_internship():
    try:
        data = request.json
        internship_domain = data.get('domain')

        if not internship_domain:
            return jsonify({"error": "Domain is required"}), 400

        # Load the generative model
        model = GenerativeModel("gemini-1.5-flash-001")
        chat = model.start_chat()

        # Generate content based on user input
        user_input = f"Create two simple internship tasks in the {internship_domain} domain with real-world code samples."
        response = chat.send_message(user_input)

        # Send Markdown response to the front-end
        return jsonify({"content": response.text})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Endpoint to download the content as .docx
@app.route('/download-docx', methods=['POST'])
def download_docx():
    try:
        data = request.json
        markdown_content = data.get('content')
        internship_domain = data.get('domain')

        # Create a new .docx document
        document = Document()
        document.add_heading(f'Internship Program Details for {internship_domain}', 0)
        document.add_paragraph(markdown_content)

        # Save the document temporarily on the server
        docx_filename = f"internship_program_{sanitize_filename(internship_domain)}.docx"
        document.save(docx_filename)

        # Schedule file deletion after sending the response
        @after_this_request
        def remove_file(response):
            try:
                os.remove(docx_filename)
            except Exception as e:
                print(f"Error deleting file: {e}")
            return response

        # Send the generated document to the user
        return send_file(docx_filename, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
